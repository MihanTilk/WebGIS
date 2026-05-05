"""Generate 226121P_doc2.docx — the WebGIS reflective technical report.

Brief, point-form prose. Heavy focus on the innovations that extend the
basic tutorial: trails, direction arrows, density heatmap, time-travel
playback, proximity search, asset interactions, dual-CRS display,
Sri Lanka land-polygon clipping, mobile responsiveness, deployment.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def heading(doc, text, size=14, bold=True, italic=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def para(doc, text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def centered(doc, text, bold=False, size=12, italic=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    return p


def blank(doc):
    doc.add_paragraph()


# ---------------------------------------------------------------------------

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


# ===== COVER PAGE ==========================================================

for _ in range(3):
    blank(doc)
centered(doc, "Department Of Decision Sciences", bold=True, size=14, after=4)
centered(doc, "Faculty Of Business", bold=True, size=14, after=4)
centered(doc, "University Of Moratuwa", bold=True, size=14, after=4)
for _ in range(2):
    blank(doc)
centered(doc, "Semester 06", bold=True, after=4)
centered(doc, "Geographic Information Systems", bold=True, after=4)
centered(doc, "WebGIS Asset Tracking Application", bold=True, after=4)
centered(doc, "Reflective Technical Report", bold=True, after=4)
for _ in range(2):
    blank(doc)
centered(doc, "M. D. Tilakaratne", bold=True, after=4)
centered(doc, "226121P", bold=True, after=4)
for _ in range(2):
    blank(doc)
centered(doc, "Date of Submission", bold=True, after=4)
centered(doc, "[06/05/2026]", bold=True, after=4)
blank(doc)
centered(doc, "Word Count", bold=True, after=4)
centered(doc, "[3185]", bold=True, after=4)
doc.add_page_break()


# ===== ACKNOWLEDGEMENT =====================================================

heading(doc, "Acknowledgement")
para(
    doc,
    "I would like to thank Dr. Sandun Dassanayake for the foundational tutorial "
    "\"Building a WebGIS Asset Tracking Application\" which served as the starting "
    "point for this project. The clear explanations of the spatial database design, "
    "REST API patterns, and frontend implementation techniques were instrumental in "
    "shaping my approach. I am also grateful for the opportunity to extend the "
    "basic tutorial in directions that allowed me to explore additional spatial "
    "concepts and practical applications."
)
doc.add_page_break()


# ===== EXECUTIVE SUMMARY ===================================================

heading(doc, "Executive Summary")
para(
    doc,
    "This report documents the development of a WebGIS asset tracking application "
    "that visualizes the locations of moving assets across Sri Lanka in near "
    "real-time. The application was built following the tutorial by Dr. Sandun "
    "Dassanayake, then extended significantly with seven major innovations beyond "
    "the basic implementation."
)
blank(doc)
para(doc, "Headline numbers:")
bullet(doc, "PostgreSQL 18 + PostGIS 3.x backend, Flask REST API, OpenLayers 7 frontend.")
bullet(doc, "14 REST API endpoints (CRUD, history, motion, snapshot, heatmap, proximity, interactions).")
bullet(doc, "~150 simulated assets across 20 Sri Lankan cities, clipped to a hand-built mainland polygon.")
bullet(doc, "Seven extension features beyond the tutorial baseline.")
bullet(doc, "Deployed publicly on Render (backend) + Supabase (database) + GitHub Pages (frontend).")
bullet(doc, "Hourly automated database cleanup via GitHub Actions to stay within free-tier limits.")
blank(doc)
para(
    doc,
    "Source code: https://github.com/MihanTilk/WebGIS"
)
doc.add_page_break()


# ===== TABLE OF CONTENTS ===================================================

heading(doc, "Table of Contents")
toc = [
    ("Acknowledgement", "i"),
    ("Executive Summary", "ii"),
    ("Table of Contents", "iii"),
    ("List of Figures", "iv"),
    ("1.0 Introduction", "1"),
    ("2.0 Aim and Objectives", "2"),
    ("3.0 System Architecture Summary", "3"),
    ("4.0 Development Procedure", "4"),
    ("    4.1 Spatial Database Design", "4"),
    ("    4.2 Flask Backend", "5"),
    ("    4.3 OpenLayers Frontend", "6"),
    ("    4.4 Simulator and Automated Cleanup", "7"),
    ("5.0 Reflection and Insights", "8"),
    ("    5.1 Innovations and How They Connect", "8"),
    ("    5.2 Challenges", "12"),
    ("    5.3 Limitations", "13"),
    ("    5.4 Improvements", "14"),
    ("    5.5 Real-World Applications", "15"),
    ("6.0 Lessons Learned", "16"),
    ("7.0 Conclusion", "17"),
    ("8.0 References", "18"),
]
for entry, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pad = max(2, 70 - len(entry) - len(page))
    p.add_run(entry + ("." * pad) + page)
doc.add_page_break()


# ===== LIST OF FIGURES =====================================================

heading(doc, "List of Figures")
for cap, page in [
    ("Figure 1: Three-tier architecture", "3"),
    ("Figure 2: Assets table schema", "4"),
    ("Figure 3: Proximity query SQL using PostGIS geography type", "5"),
    ("Figure 4: Periodic refresh loop", "6"),
    ("Figure 5: Innovation feature map and shared infrastructure", "11"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pad = max(2, 80 - len(cap) - len(page))
    p.add_run(cap + ("." * pad) + page)
doc.add_page_break()


# ===== 1.0 INTRODUCTION ====================================================

heading(doc, "1.0 Introduction")
para(
    doc,
    "WebGIS (Web-based Geographic Information System) refers to applications that "
    "deliver spatial data and geographic visualization through web browsers. WebGIS "
    "for asset tracking specifically stores the locations of physical assets in a "
    "spatial database, exposes them through web APIs, and renders them on an "
    "interactive map that updates as the assets move."
)
blank(doc)
para(doc, "The application discussed in this report:")
bullet(doc, "Was built following the tutorial \"Building a WebGIS Asset Tracking Application\" by Dr. Sandun Dassanayake (2025).")
bullet(doc, "Uses PostgreSQL with PostGIS for spatial storage, Flask for the REST API, and OpenLayers for the interactive map.")
bullet(doc, "Was extended significantly beyond the basic tutorial with seven innovation features described in Section 5.1.")
bullet(doc, "Is deployed publicly with the source available at https://github.com/MihanTilk/WebGIS.")
blank(doc)
para(doc, "Motivations for the project:")
bullet(doc, "Hands-on experience with spatial database design, REST API development, projection handling, and interactive map programming.")
bullet(doc, "Direct relevance to real-world domains such as logistics, fleet management, field operations, and emergency response.")
bullet(doc, "Custom-built tracking solutions allow organizations to retain control of their data instead of relying on third-party platforms.")
doc.add_page_break()


# ===== 2.0 AIM AND OBJECTIVES ==============================================

heading(doc, "2.0 Aim and Objectives")
heading(doc, "Aim:", size=12, italic=True, bold=False, before=6)
para(
    doc,
    "To design and implement a WebGIS-based asset tracking application that "
    "demonstrates the integration of PostGIS, Flask, and OpenLayers, and to "
    "critically reflect on the development process, design decisions, and "
    "innovation features added beyond the tutorial."
)
heading(doc, "Objectives:", size=12, italic=True, bold=False, before=6)
numbered(doc, "Design a spatial database in PostgreSQL with PostGIS, storing asset locations as Point geometry in WGS84 (EPSG:4326) with a GIST spatial index, by 06 May 2026.")
numbered(doc, "Implement a Flask REST API that returns asset data as standard GeoJSON, supporting at least 10 endpoints covering CRUD plus spatial queries.")
numbered(doc, "Develop an OpenLayers frontend that displays at least 100 simulated moving assets across Sri Lanka with periodic 5-second polling.")
numbered(doc, "Build a Python simulator that updates asset positions and triggers interaction event detection on a configurable interval.")
numbered(doc, "Extend the basic tutorial with at least 4 additional features that demonstrate deeper understanding of the underlying spatial technologies.")
numbered(doc, "Deploy the application publicly with automated database housekeeping to stay within free-tier service limits.")
numbered(doc, "Document the development process and reflect on design decisions, limitations, and possible improvements.")
doc.add_page_break()


# ===== 3.0 SYSTEM ARCHITECTURE =============================================

heading(doc, "3.0 System Architecture Summary")
para(
    doc,
    "The system follows the standard WebGIS three-tier architecture. Each layer is "
    "developed and deployed independently. A separate Python simulator runs on a "
    "developer laptop or laptop-during-demo basis."
)
blank(doc)
centered(doc, "OpenLayers (browser)", bold=True, after=2)
centered(doc, "|", after=2)
centered(doc, "HTTP / GeoJSON", italic=True, size=10, after=2)
centered(doc, "|", after=2)
centered(doc, "Flask (Python REST API)", bold=True, after=2)
centered(doc, "|", after=2)
centered(doc, "SQL (PostGIS queries)", italic=True, size=10, after=2)
centered(doc, "|", after=2)
centered(doc, "PostgreSQL + PostGIS", bold=True, after=8)
centered(doc, "Figure 1: Three-tier architecture", italic=True, size=10, after=12)
para(doc, "Layer responsibilities:")
bullet(doc, "Database: PostgreSQL 18 + PostGIS 3.x. Three tables (assets, asset_history, interactions) with GIST indexes on every geometry column.")
bullet(doc, "Backend: Flask + psycopg2 + flask-cors. 14 REST endpoints returning GeoJSON.")
bullet(doc, "Frontend: OpenLayers 7 over OpenStreetMap tiles. Polling with setInterval; vector layers for markers, trails, arrows, encounters, heatmap, proximity circle.")
bullet(doc, "Simulator: standalone Python script that PUTs new positions and triggers interaction detection on a configurable interval.")
blank(doc)
para(doc, "Deployment topology:")
bullet(doc, "Backend hosted on Render (free web service, sleeps after 15 minutes idle).")
bullet(doc, "Database hosted on Supabase (free tier, 500 MB cap).")
bullet(doc, "Frontend served from GitHub Pages.")
bullet(doc, "Hourly cleanup runs as a GitHub Actions workflow, keeping asset_history bounded to a rolling 2-hour window.")
doc.add_page_break()


# ===== 4.0 DEVELOPMENT PROCEDURE ===========================================

heading(doc, "4.0 Development Procedure")
para(doc, "Five development stages, each completed and tested before moving to the next.")

heading(doc, "4.1 Spatial Database Design", size=13, before=12)
para(doc, "Three tables, each with a clear single responsibility:")
bullet(doc, "assets: current state. One row per asset, Point geometry in WGS84, GIST spatial index.")
bullet(doc, "asset_history: every position change. Populated by an AFTER INSERT/UPDATE trigger on assets.geom.")
bullet(doc, "interactions: typed proximity encounters between asset pairs (5 kinds). Asset_a < asset_b check constraint to deduplicate.")
code(
    doc,
    "CREATE TABLE assets (\n"
    "    id SERIAL PRIMARY KEY,\n"
    "    name TEXT NOT NULL,\n"
    "    asset_type TEXT NOT NULL,\n"
    "    last_seen TIMESTAMP DEFAULT NOW(),\n"
    "    geom GEOMETRY(Point, 4326) NOT NULL\n"
    ");\n"
    "CREATE INDEX idx_assets_geom ON assets USING GIST (geom);"
)
centered(doc, "Figure 2: Assets table schema", italic=True, size=10, after=12)
para(doc, "Two seed scripts:")
bullet(doc, "02_sample_data.sql: 6 assets around Colombo for the basic demo.")
bullet(doc, "04_clustered_seed.sql: ~150 assets across 20 Sri Lankan cities. Two tiers (clustered + roaming), both clipped to a hand-built mainland polygon via ST_Contains.")
bullet(doc, "Includes a synthetic past-walk backfill so trail visualization works immediately without needing the simulator to accumulate data first.")
bullet(doc, "Includes a synthetic interactions backfill so the side panel shows events on first load.")

heading(doc, "4.2 Flask Backend", size=13, before=12)
para(doc, "14 REST endpoints. Selected highlights:")
bullet(doc, "GET /api/assets, /api/assets/<id> — read.")
bullet(doc, "POST /api/assets — create. PUT /api/assets/<id> — update location. DELETE — remove.")
bullet(doc, "GET /api/assets/<id>/history — recent positions as GeoJSON LineString trail.")
bullet(doc, "GET /api/assets/<id>/motion — speed (km/h) and heading (degrees) from last two history rows.")
bullet(doc, "GET /api/snapshot?at=<iso> — asset positions at any past timestamp (powers playback).")
bullet(doc, "GET /api/heatmap?hours=N — recency-weighted point cloud for kernel density visualization.")
bullet(doc, "GET /api/proximity?lon&lat&radius_m — assets within a given metric radius.")
bullet(doc, "GET /api/interactions, POST /api/interactions/detect, GET /api/interactions/rules — typed encounter system.")
bullet(doc, "GET /api/health — DB-independent liveness probe for Render's health check.")
para(doc, "All distance work uses the PostGIS geography type so radii are in real metres on the WGS84 spheroid:")
code(
    doc,
    "SELECT id, name, asset_type, ST_AsGeoJSON(geom) AS geom_json,\n"
    "       ST_Distance(geom::geography,\n"
    "                   ST_SetSRID(ST_MakePoint(%s, %s), 4326)::geography)\n"
    "           AS distance_m\n"
    "FROM assets\n"
    "WHERE ST_DWithin(geom::geography,\n"
    "                 ST_SetSRID(ST_MakePoint(%s, %s), 4326)::geography, %s)\n"
    "ORDER BY distance_m ASC;"
)
centered(doc, "Figure 3: Proximity query (PostGIS geography type)", italic=True, size=10, after=12)

heading(doc, "4.3 OpenLayers Frontend", size=13, before=12)
para(doc, "Single self-contained HTML file. Layer stack (bottom to top):")
bullet(doc, "OpenStreetMap raster tiles.")
bullet(doc, "Heatmap layer (recency-weighted KDE, zoom-aware kernel sizing).")
bullet(doc, "Proximity search circle.")
bullet(doc, "Encounter circle (when an interaction event is selected).")
bullet(doc, "Movement trail (dashed coloured line per clicked asset).")
bullet(doc, "Direction arrow (rotated SVG, yellow when speed >= 5 km/h).")
bullet(doc, "Interaction stars at event midpoints.")
bullet(doc, "Asset markers (always on top to remain clickable).")
para(doc, "Periodic refresh:")
code(
    doc,
    "async function refresh() {\n"
    "    const res = await fetch(API_BASE, { cache: 'no-store' });\n"
    "    const data = await res.json();\n"
    "    const features = new ol.format.GeoJSON().readFeatures(data, {\n"
    "        featureProjection: 'EPSG:3857'\n"
    "    });\n"
    "    vectorSource.clear();\n"
    "    vectorSource.addFeatures(features);\n"
    "}\n"
    "setInterval(refresh, REFRESH_MS);"
)
centered(doc, "Figure 4: Periodic refresh loop", italic=True, size=10, after=12)
para(doc, "UI controls:")
bullet(doc, "Type filter (vehicle / person / equipment / all).")
bullet(doc, "Live vs. Playback mode toggle, with time slider, play button, and speed selector (1x / 5x / 30x).")
bullet(doc, "Density heatmap toggle with hours selector (1 / 6 / 24 / 72).")
bullet(doc, "Proximity search toggle with radius slider (100 m to 50 km).")
bullet(doc, "Fit-to-all button.")
bullet(doc, "Mobile-responsive layout: panels collapse, controls reflow, compass repositions below zoom buttons.")

heading(doc, "4.4 Simulator and Automated Cleanup", size=13, before=12)
para(doc, "The simulator (simulator.py) drives the system without real GPS data:")
bullet(doc, "Random-walks every existing asset on a configurable tick (default every 2 s).")
bullet(doc, "Optionally auto-spawns new assets, scattered around 20 Sri Lankan cities.")
bullet(doc, "After each tick, calls POST /api/interactions/detect to advance the encounter detection.")
bullet(doc, "CLI flags: --interval, --jitter, --island, --spawn-every, --spawn-max.")
para(doc, "Automated database housekeeping:")
bullet(doc, "GitHub Actions workflow runs every hour (.github/workflows/cleanup.yml).")
bullet(doc, "Deletes asset_history rows older than 2 hours and closed interactions older than 2 hours.")
bullet(doc, "Runs VACUUM ANALYZE to reclaim disk space.")
bullet(doc, "Keeps the cloud database bounded at roughly 30 MB regardless of how long the simulator runs.")
para(doc, "Smoke test:")
bullet(doc, "PowerShell script (smoke_test.ps1) exercises every endpoint plus four error paths.")
bullet(doc, "Creates and deletes a test asset so the database is left in its original state.")

doc.add_page_break()


# ===== 5.0 REFLECTION AND INSIGHTS =========================================

heading(doc, "5.0 Reflection and Insights")

# ----- 5.1 Innovations -----------------------------------------------------

heading(doc, "5.1 Innovations and How They Connect", size=13, before=8)
para(
    doc,
    "Seven extension features were added beyond the tutorial baseline. Each one is "
    "described below, followed by a summary of how they share infrastructure."
)

heading(doc, "Movement trails", size=12, italic=True, bold=False, before=10)
bullet(doc, "Per-asset trail, drawn as a dashed coloured polyline when an asset is clicked.")
bullet(doc, "Backed by the asset_history table, populated automatically by an AFTER INSERT/UPDATE trigger on assets.geom.")
bullet(doc, "GET /api/assets/<id>/history returns recent positions ordered ASC for line drawing.")
bullet(doc, "Synthetic past-walk in the seed script ensures trails are visible immediately, not just after the simulator runs.")

heading(doc, "Direction arrow with speed and heading", size=12, italic=True, bold=False, before=10)
bullet(doc, "Inline SVG arrow icon at the asset's current position, rotated to point in its direction of travel.")
bullet(doc, "Rotation in radians from PostGIS ST_Azimuth applied between the last two history rows.")
bullet(doc, "Speed computed as ST_Distance(geography) / time delta, displayed in km/h.")
bullet(doc, "Arrow turns yellow when speed >= 5 km/h, white otherwise.")

heading(doc, "Density heatmap (zoom-aware)", size=12, italic=True, bold=False, before=10)
bullet(doc, "OpenLayers Heatmap layer fed by history rows, weighted by recency in [0, 1].")
bullet(doc, "Per-point weight intentionally low (0.10x of backend value) so density emerges from accumulation rather than each point saturating the gradient.")
bullet(doc, "Custom gradient (transparent to blue to cyan to green to yellow to red) so empty areas remain clear.")
bullet(doc, "Kernel radius and blur scale linearly with zoom: tight at island view, loose at city view.")
bullet(doc, "Recomputes on zoom change with an 80 ms debounce.")

heading(doc, "Time-travel playback", size=12, italic=True, bold=False, before=10)
bullet(doc, "Mode toggle (Live / Playback). Playback uses GET /api/snapshot?at=<iso>.")
bullet(doc, "GET /api/history/range provides the slider's earliest and latest bounds.")
bullet(doc, "Slider scrubs through history; play button auto-advances at 1x / 5x / 30x speed.")
bullet(doc, "In-flight snapshot requests deduplicated with a Symbol-based ticket pattern (avoids race conditions when scrubbing fast).")
bullet(doc, "Switching back to Live mode resumes the periodic polling loop.")

heading(doc, "Proximity search", size=12, italic=True, bold=False, before=10)
bullet(doc, "Click-to-search workflow: enable proximity mode, click anywhere on the map, see the search circle and matching assets.")
bullet(doc, "GET /api/proximity uses PostGIS ST_DWithin on the geography type so the radius is true metres on the WGS84 spheroid.")
bullet(doc, "Search circle drawn as a polygon approximation of a true geodesic circle to avoid Web Mercator latitude distortion.")
bullet(doc, "Matched assets get a yellow halo and remain full-colour; non-matches fade to grey 35% opacity.")
bullet(doc, "Radius slider re-runs the query live as it moves.")

heading(doc, "Asset interaction events", size=12, italic=True, bold=False, before=10)
bullet(doc, "Five typed proximity rules: PICKUP (vehicle x person), LOADING (vehicle x equipment), OPERATING (person x equipment), MEETING (person x person), CONVOY (vehicle x vehicle).")
bullet(doc, "Detection is a spatial join on the assets table: ST_DWithin on the geography type, joined with the rules table on type-pair match.")
bullet(doc, "Idempotent open / close logic. Inserts on first proximity, updates ended_at when the pair drifts out of range.")
bullet(doc, "Side panel polls /api/interactions every 5 s, renders filter chips, allows click-to-zoom-and-highlight on each event.")
bullet(doc, "Each event marker is a 5-pointed star at the encounter midpoint, color-coded by kind.")
bullet(doc, "Clicking an event in the side panel: highlights the two parties with a yellow halo, draws an encounter circle sized to the rule's proximity threshold, and zooms to fit both.")

heading(doc, "Sri Lanka Kandawala Grid coordinate display", size=12, italic=True, bold=False, before=10)
bullet(doc, "Popup shows asset position in WGS84 (EPSG:4326) AND Sri Lanka Kandawala Grid (EPSG:5234, Transverse Mercator).")
bullet(doc, "Backend uses ST_Transform(geom, 5234) to compute easting / northing in the same query that returns the GeoJSON geometry.")
bullet(doc, "Domain-relevant: Sri Lankan civil engineering and cadastral work uses EPSG:5234 metres, not WGS84 degrees.")

heading(doc, "Land-polygon clipping", size=12, italic=True, bold=False, before=10)
bullet(doc, "Hand-built ~38-vertex polygon of the Sri Lanka mainland, applied via ST_Contains.")
bullet(doc, "Used by both the clustered seed (over-generates 12 candidates per city, takes first 6 on land) and the roaming tier (over-generates 200, takes first 30 on land).")
bullet(doc, "Same containment predicate as a geofencing system, used here as a generative filter.")

heading(doc, "How the innovations connect", size=12, italic=True, bold=False, before=10)
para(
    doc,
    "Six of the seven innovations share two pieces of infrastructure: the "
    "asset_history table and the geography-type spatial functions. The history "
    "table is populated automatically by a single trigger on assets.geom and is "
    "the source of truth for trails, direction arrows (last two rows), playback "
    "(snapshot at a past timestamp), heatmap (full window of recent rows), and "
    "interaction event detection (paired with current positions). The geography "
    "type ensures every distance, area, and proximity comparison is in real "
    "metres on the spheroid: ST_DWithin for proximity and interactions, "
    "ST_Distance for proximity ranking and motion speed, ST_Azimuth for heading."
)
blank(doc)
centered(doc, "asset_history (event-sourced trigger)", bold=True, after=4)
centered(doc, "feeds:", italic=True, size=10, after=4)
centered(doc, "trails | arrows | heatmap | playback | interactions", after=8)
centered(doc, "geography type (real metres on WGS84 spheroid)", bold=True, after=4)
centered(doc, "powers:", italic=True, size=10, after=4)
centered(doc, "proximity | interactions | speed/heading", after=8)
centered(doc, "ST_Contains (point-in-polygon)", bold=True, after=4)
centered(doc, "powers:", italic=True, size=10, after=4)
centered(doc, "land-clipping for both seed tiers", after=8)
centered(doc, "Figure 5: Innovation feature map and shared infrastructure",
         italic=True, size=10, after=12)
para(
    doc,
    "This pattern (one well-designed schema element supporting many features) is "
    "the core lesson of the innovation work. The history table and the geography "
    "type were each added as a single small piece of infrastructure but unlocked "
    "five user-facing capabilities each. The same applies to ST_Contains: added "
    "for land-clipping, but the same predicate would also drive a geofence layer "
    "if added later."
)

# ----- 5.2 Challenges ------------------------------------------------------

heading(doc, "5.2 Challenges", size=13, before=12)
bullet(doc, "Timezone handling: TIMESTAMP without time zone in PostgreSQL vs. UTC ISO strings from the frontend caused snapshot queries to silently return empty results until the parsing was rewritten to be UTC-aware.")
bullet(doc, "Interaction thresholds: real-world values (5 to 20 m, 10 to 60 s) never fired against the random-walk simulator. Thresholds were loosened to demo values (80 to 300 m, 4 to 6 s); real values preserved as in-line comments.")
bullet(doc, "Sri Lanka land polygon: initial 19-vertex outline cut chords through bays. Refined to 38 vertices following the actual coastline.")
bullet(doc, "Heatmap saturation: per-point weights too high, so single isolated assets painted full-red blobs. Fixed by reducing per-point weight to 10% and adding zoom-aware kernel sizing.")
bullet(doc, "Mobile compass overlap: the custom compass control was repositioned at top-left on mobile and overlapped the OpenLayers zoom buttons. Moved below the zoom column.")
bullet(doc, "CORS errors during local development, surfacing as generic network failures. Fixed with flask-cors.")
bullet(doc, "Connection leak in early backend code (with conn pattern doesn't close). Refactored to try/finally with explicit close.")
bullet(doc, "XSS risk in popup HTML template (innerHTML with asset name). Fixed by using textContent.")
bullet(doc, "history endpoint LIMIT was returning oldest N within window instead of most recent N. Fixed by wrapping in a DESC subquery before re-ordering ASC.")

# ----- 5.3 Limitations -----------------------------------------------------

heading(doc, "5.3 Limitations", size=13, before=12)
bullet(doc, "Periodic polling, not push. Asset positions can be up to 5 seconds out of date.")
bullet(doc, "Simulator uses random walk, which does not match real GPS movement statistics.")
bullet(doc, "No connection pooling. Each request opens a new psycopg2 connection.")
bullet(doc, "No authentication or authorization. The API is fully open.")
bullet(doc, "Sri Lanka land polygon is hand-built (~38 vertices). Excludes Mannar Island, the Jaffna islands, and small offshore landforms.")
bullet(doc, "asset_history grows unboundedly without the cleanup workflow. Free-tier Supabase (500 MB) would fill in roughly 18 hours at 2-second simulator ticks.")
bullet(doc, "Render free tier sleeps after 15 minutes idle, causing a 30-second cold start on the first request after a quiet period.")
bullet(doc, "Flask development server with debug mode enabled. Acceptable for local and demo use, not for production.")

# ----- 5.4 Improvements ----------------------------------------------------

heading(doc, "5.4 Improvements", size=13, before=12)
bullet(doc, "Real-time updates via WebSockets or Server-Sent Events instead of 5-second polling.")
bullet(doc, "JWT-based authentication and role-based access control. Tenant isolation per organization.")
bullet(doc, "Geofencing: zones table and zone_events table populated by a trigger using ST_Contains. Frontend draw tool for users to define zones.")
bullet(doc, "Mobile companion app that sends real device GPS via the existing PUT endpoint, replacing the simulator with real data.")
bullet(doc, "History retention policy: aggregate rows older than 30 days into hourly summaries, archive raw rows to cold storage.")
bullet(doc, "Production deployment: gunicorn + nginx, Supabase Pro for the database, Cloudflare Pages for the frontend, environment variables managed through a secrets manager.")
bullet(doc, "Analytics dashboards: average daily distance per asset, common routes, dwell time per zone, predicted arrival times.")
bullet(doc, "Import the official Sri Lanka boundary GeoJSON from the Survey Department for an exact land mask.")

# ----- 5.5 Real-world applications ----------------------------------------

heading(doc, "5.5 Real-World Applications", size=13, before=12)
bullet(doc, "Logistics and fleet management: dispatcher uses proximity search to assign nearest driver; trail visualization verifies route adherence; interaction events log pickups and drop-offs automatically.")
bullet(doc, "Field operations and surveying: Sri Lanka Kandawala Grid display matches the projection used by Sri Lankan civil engineering practice; heatmap surfaces under-attended sites.")
bullet(doc, "Campus management: track shuttle buses, security personnel, maintenance vehicles; students consult the live map for shuttle ETAs.")
bullet(doc, "Emergency response: real-time tracking of all responders during incidents; proximity search identifies units within an evacuation zone; interaction events flag responder rendezvous at staging areas.")
bullet(doc, "Workplace safety and contact tracing: persistent close-proximity events flagged for review; cluster source identification during disease outbreaks.")
bullet(doc, "Conservation: GPS trackers on endangered animals or valuable equipment; geofencing extension alerts rangers when tracked subjects enter dangerous areas or leave protected zones.")
doc.add_page_break()


# ===== 6.0 LESSONS LEARNED =================================================

heading(doc, "6.0 Lessons Learned")
para(doc, "Five lessons from the development process:")
bullet(doc, "Let the database do spatial work. PostGIS functions (ST_DWithin, ST_Distance, ST_Azimuth, ST_Transform) are faster and more correct than equivalent Python implementations, and the GIST index makes them fast at any data size.")
bullet(doc, "Test against realistic data. The interaction detection feature was specified for tight real-world thresholds and worked correctly on paper, but never fired against the random-walk simulator. Synthetic test data with different statistical properties from production data can hide design flaws.")
bullet(doc, "Separate current state from historical record. A single trigger that populates asset_history on every position change unlocked five distinct features (trails, motion arrows, playback, heatmap, interaction events) without any additional schema work.")
bullet(doc, "Document trade-offs in the source code. Loosened interaction thresholds for the demo, but kept the original real-world values as in-line comments so the design intent is not lost.")
bullet(doc, "Work in small increments. The smoke test script makes regression detection cheap; bugs caught immediately are bugs that do not propagate.")
doc.add_page_break()


# ===== 7.0 CONCLUSION ======================================================

heading(doc, "7.0 Conclusion")
para(
    doc,
    "This report has documented the development of a WebGIS-based asset tracking "
    "application that integrates PostgreSQL with PostGIS, Flask, and OpenLayers in "
    "a three-tier architecture. The application meets and exceeds the requirements "
    "of the assignment, simulating roughly 150 moving assets across Sri Lanka, "
    "storing them spatially, updating periodically, and visualising them on an "
    "interactive map with full metadata."
)
para(doc, "Aim achieved:")
bullet(doc, "14 REST endpoints implemented covering CRUD plus seven categories of spatial query.")
bullet(doc, "Seven innovation features beyond the tutorial baseline: trails, direction arrows, density heatmap, time-travel playback, proximity search, asset interactions, dual-CRS coordinate display.")
bullet(doc, "Mobile-responsive layout with collapsible panels and repositioned controls.")
bullet(doc, "Public deployment with hourly automated database cleanup keeping the system within free-tier limits indefinitely.")
para(
    doc,
    "Limitations identified (polling rather than push, random-walk simulator, no "
    "authentication, simplified land polygon) are matched by clear improvement "
    "paths described in Section 5.4. The skills developed are directly transferable "
    "to logistics, field operations, campus management, and emergency response."
)
para(doc, "Source code: https://github.com/MihanTilk/WebGIS")
doc.add_page_break()


# ===== 8.0 REFERENCES ======================================================

heading(doc, "8.0 References")
refs = [
    "Dassanayake, S. (2025). Building a WebGIS Asset Tracking Application: A "
    "Comprehensive Tutorial. University of Moratuwa.",

    "PostGIS Project Steering Committee. (2024). PostGIS Documentation. Retrieved "
    "from https://postgis.net/documentation/",

    "Pallets. (2024). Flask Documentation. Retrieved from "
    "https://flask.palletsprojects.com/",

    "OpenLayers Contributors. (2024). OpenLayers API Documentation. Retrieved from "
    "https://openlayers.org/en/latest/doc/",

    "Open Source Geospatial Foundation. (2024). OpenStreetMap Tile Usage Policy. "
    "Retrieved from https://operations.osmfoundation.org/policies/tiles/",

    "Butler, H., Daly, M., Doyle, A., Gillies, S., Hagen, S., & Schaub, T. (2016). "
    "The GeoJSON Format (RFC 7946). Internet Engineering Task Force. Retrieved "
    "from https://geojson.org/",

    "International Association of Oil and Gas Producers. (2024). EPSG Geodetic "
    "Parameter Dataset. Retrieved from https://epsg.org/",

    "Open Geospatial Consortium. (2008). ISO 19125-1:2004 Geographic information "
    "- Simple feature access - Part 1: Common architecture.",

    "Survey Department of Sri Lanka. (n.d.). Sri Lanka Datum and Projection "
    "Systems. Retrieved from https://www.survey.gov.lk/",

    "Render, Inc. (2024). Render Free Tier Documentation. Retrieved from "
    "https://render.com/docs/free",

    "Supabase. (2024). Supabase PostgreSQL Documentation. Retrieved from "
    "https://supabase.com/docs/guides/database",

    "Tilakaratne, M. D. (2026). WebGIS Asset Tracking Application [Source code]. "
    "GitHub. Retrieved from https://github.com/MihanTilk/WebGIS",
]
for r in refs:
    p = doc.add_paragraph(r, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


import os

out = "c:/Users/mitil/OneDrive/Desktop/226121P/226121P_doc2.docx"
try:
    doc.save(out)
    print(f"Saved: {out}")
except PermissionError:
    fallback = "c:/Users/mitil/OneDrive/Desktop/226121P/226121P_doc2_NEW.docx"
    doc.save(fallback)
    print(f"Original locked (probably open in Word). Saved instead to: {fallback}")
    print("Close Word, delete 226121P_doc2.docx, rename _NEW.docx, and you're done.")
